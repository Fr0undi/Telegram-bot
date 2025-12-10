"""Сервис форматирования документов по ГОСТ"""

import os
import re
import traceback

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from app.config.settings import settings
from app.utils.logger import setup_logger


logger = setup_logger(__name__)


def _find_title_page_end(doc: Document) -> int:
    """
    Находит индекс параграфа, где заканчивается титульный лист.
    Титульник заканчивается перед "Содержание" или "Оглавление".

    Returns:
        Индекс первого параграфа после титульника
    """
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip().upper()
        if text in ['СОДЕРЖАНИЕ', 'ОГЛАВЛЕНИЕ']:
            return i
    return 0


def _is_main_heading(text: str) -> bool:
    """
    Проверяет, является ли текст главным заголовком (требует разрыв страницы).
    """
    text_upper = text.strip().upper()

    # Специальные разделы
    if text_upper in ['СОДЕРЖАНИЕ', 'ОГЛАВЛЕНИЕ', 'ВВЕДЕНИЕ', 'ЗАКЛЮЧЕНИЕ',
                      'СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ', 'БИБЛИОГРАФИЧЕСКИЙ СПИСОК',
                      'АННОТАЦИЯ', 'РЕФЕРАТ', 'ПРИЛОЖЕНИЕ', 'ПРИЛОЖЕНИЯ']:
        return True

    # Главы
    if re.match(r'^ГЛАВА\s+\d+', text_upper):
        return True

    return False


def _is_subheading(text: str) -> bool:
    """
    Проверяет, является ли текст подзаголовком (1.1, 1.2, 2.1 и т.д.)
    """
    return bool(re.match(r'^\d+\.\d+', text.strip()))


def _is_figure_caption(text: str) -> bool:
    """Проверяет, является ли текст подписью к рисунку"""
    return bool(re.match(r'^Рисунок\s*\d*\s*[-–—]', text.strip(), re.IGNORECASE))


def _is_table_caption(text: str) -> bool:
    """Проверяет, является ли текст подписью к таблице"""
    return bool(re.match(r'^Таблица\s+\d+\s*[-–—]', text.strip(), re.IGNORECASE))


def _has_image(paragraph) -> bool:
    """Проверяет, содержит ли параграф изображение."""
    # Ищем blip элементы (изображения)
    blip = paragraph._element.find('.//{http://schemas.openxmlformats.org/drawingml/2006/main}blip')
    if blip is not None:
        return True

    # Также проверяем drawing элементы
    drawing = paragraph._element.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing')
    if drawing is not None:
        return True

    # И pict элементы (старый формат)
    pict = paragraph._element.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pict')
    if pict is not None:
        return True

    return False


def _is_list_item(paragraph) -> bool:
    """Проверяет, является ли параграф элементом списка (имеет numPr)"""
    pPr = paragraph._element.find(qn('w:pPr'))
    if pPr is not None:
        numPr = pPr.find(qn('w:numPr'))
        return numPr is not None
    return False


def _fix_list_first_letter(doc: Document, title_end_idx: int) -> None:
    """
    Делает первую букву в элементах списка строчной.
    """
    count = 0

    for i, para in enumerate(doc.paragraphs):
        if i < title_end_idx:
            continue

        if not _is_list_item(para):
            continue

        if not para.runs:
            continue

        # Находим первый непустой run
        for run in para.runs:
            text = run.text
            if not text or not text.strip():
                continue

            # Находим первую букву
            first_letter_idx = None
            for j, char in enumerate(text):
                if char.isalpha():
                    first_letter_idx = j
                    break

            if first_letter_idx is not None and text[first_letter_idx].isupper():
                # Проверяем что это не аббревиатура (следующая буква тоже заглавная)
                next_idx = first_letter_idx + 1
                if next_idx < len(text) and text[next_idx].isupper():
                    break  # Это аббревиатура, не трогаем

                # Делаем первую букву строчной
                new_text = text[:first_letter_idx] + text[first_letter_idx].lower() + text[first_letter_idx+1:]
                run.text = new_text
                count += 1
            break

    logger.info(f"Исправлено первых букв в списках: {count}")


def _fix_multiple_spaces(doc: Document, title_end_idx: int) -> None:
    """
    Убирает множественные пробелы (заменяет на один).
    Обрабатывает случаи когда пробелы в разных runs.
    НЕ трогает титульник.
    """
    import re
    count = 0

    def process_paragraph(para):
        nonlocal count
        runs = para.runs

        # Сначала обрабатываем внутри каждого run
        for run in runs:
            if not run.text:
                continue

            original = run.text
            new_text = re.sub(r' {2,}', ' ', original)

            if new_text != original:
                run.text = new_text
                count += 1

        # Затем обрабатываем пробелы между runs
        for i in range(len(runs) - 1, 0, -1):  # Идём с конца
            curr_run = runs[i]
            prev_run = runs[i - 1]

            if not curr_run.text or not prev_run.text:
                continue

            # Если предыдущий run заканчивается на пробел, а текущий начинается с пробела
            if prev_run.text.endswith(' ') and curr_run.text.startswith(' '):
                # Убираем пробел в начале текущего run
                curr_run.text = curr_run.text.lstrip(' ')
                count += 1

            # Если текущий run содержит только пробелы
            if curr_run.text.strip() == '' and prev_run.text.endswith(' '):
                curr_run.text = ''
                count += 1

    for i, para in enumerate(doc.paragraphs):
        if i < title_end_idx:
            continue
        process_paragraph(para)

    # Также в таблицах (они не в титульнике)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    process_paragraph(para)

    logger.info(f"Убрано множественных пробелов в {count} местах")


def _fix_abbreviations(doc: Document, title_end_idx: int) -> None:
    """
    Заменяет сокращения на полные формы:
    - т.к. → так как
    - т.е. → то есть
    - и т.д. → и так далее
    - и т.п. → и тому подобное
    - и др. → и другие

    НЕ заменяет:
    - т.н. (слишком часто путается с инициалами Т. Н.)

    НЕ трогает титульник.
    """
    import re
    count = 0

    def process_text(text):
        if not text:
            return text, False

        new_text = text

        # Важно: используем отрицательный lookbehind чтобы не заменять после запятой (инициалы)
        # т.к. → так как (но не после запятой, как в "Иванов, т.к. ...")
        new_text = re.sub(r'(?<![,А-ЯA-Z])\s+т\.\s*к\.', ' так как', new_text, flags=re.IGNORECASE)
        new_text = re.sub(r'^т\.\s*к\.', 'так как', new_text, flags=re.IGNORECASE)

        # т.е. → то есть
        new_text = re.sub(r'(?<![,А-ЯA-Z])\s+т\.\s*е\.', ' то есть', new_text, flags=re.IGNORECASE)
        new_text = re.sub(r'^т\.\s*е\.', 'то есть', new_text, flags=re.IGNORECASE)

        # и т.д. → и так далее
        new_text = re.sub(r'\bи\s+т\.\s*д\.', 'и так далее', new_text, flags=re.IGNORECASE)

        # и т.п. → и тому подобное
        new_text = re.sub(r'\bи\s+т\.\s*п\.', 'и тому подобное', new_text, flags=re.IGNORECASE)

        # и др. → и другие
        new_text = re.sub(r'\bи\s+др\.', 'и другие', new_text, flags=re.IGNORECASE)

        return new_text, new_text != text

    for i, para in enumerate(doc.paragraphs):
        if i < title_end_idx:
            continue
        for run in para.runs:
            if run.text:
                new_text, changed = process_text(run.text)
                if changed:
                    run.text = new_text
                    count += 1

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        if run.text:
                            new_text, changed = process_text(run.text)
                            if changed:
                                run.text = new_text
                                count += 1

    logger.info(f"Заменено сокращений в {count} runs")


def _format_bibliography(doc: Document) -> None:
    """
    Форматирует библиографический список:
    - Шрифт Times New Roman 14pt
    - Выравнивание по ширине
    - Не меняет окончания записей
    """
    count = 0
    in_bibliography = False

    for para in doc.paragraphs:
        text = para.text.strip()

        # Определяем начало библиографии
        if 'БИБЛИОГРАФИЧЕСКИЙ' in text.upper() or 'СПИСОК ЛИТЕРАТУРЫ' in text.upper():
            in_bibliography = True
            continue

        # Определяем конец библиографии
        if in_bibliography:
            if text.upper().startswith('ПРИЛОЖЕНИЕ') or text.upper().startswith('ГЛАВА'):
                in_bibliography = False
                continue

            # Форматируем элементы библиографии
            if text and not text.upper() == text:  # Не заголовок
                para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                para.paragraph_format.first_line_indent = Cm(0)
                para.paragraph_format.left_indent = Cm(0)
                para.paragraph_format.line_spacing = settings.GOST_LINE_SPACING

                # Шрифт
                for run in para.runs:
                    run.font.name = settings.GOST_FONT_NAME
                    run.font.size = Pt(settings.GOST_FONT_SIZE)
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), settings.GOST_FONT_NAME)

                count += 1

    logger.info(f"Отформатировано {count} записей в библиографии")


def _add_table_captions(doc: Document) -> None:
    """
    Проверяет и форматирует подписи к таблицам.
    Формат: "Таблица X – Название"
    """
    import re
    count = 0

    for para in doc.paragraphs:
        text = para.text.strip()

        # Ищем подписи таблиц
        if text.lower().startswith('таблица'):
            # Проверяем формат: Таблица N – Название
            match = re.match(r'^[Тт]аблица\s*(\d+)\s*[-–—]\s*(.+)$', text)

            if match:
                num = match.group(1)
                title = match.group(2)

                # Исправляем формат (короткое тире)
                new_text = f"Таблица {num} – {title}"

                if new_text != text:
                    para.clear()
                    run = para.add_run(new_text)
                    run.font.name = settings.GOST_FONT_NAME
                    run.font.size = Pt(settings.GOST_FONT_SIZE)
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), settings.GOST_FONT_NAME)
                    count += 1

            # Выравнивание по центру для подписей таблиц
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            para.paragraph_format.first_line_indent = Cm(0)

    if count > 0:
        logger.info(f"Исправлено {count} подписей таблиц")
    else:
        logger.info("Подписи таблиц не найдены или уже корректны")


def _format_tables(doc: Document) -> None:
    """
    Форматирует таблицы:
    - Шрифт Times New Roman 12pt
    - Выравнивание текста
    """
    count = 0

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    para.alignment = WD_ALIGN_PARAGRAPH.LEFT

                    for run in para.runs:
                        run.font.name = settings.GOST_FONT_NAME
                        run.font.size = Pt(12)
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), settings.GOST_FONT_NAME)
                        count += 1

    if count > 0:
        logger.info(f"Отформатировано {count} runs в таблицах")
    else:
        logger.info("Таблицы не найдены или пусты")


def _fix_non_breaking_spaces(doc: Document, title_end_idx: int) -> None:
    """
    Добавляет неразрывные пробелы:
    - После №, §, рис., табл.
    - Между инициалами и фамилией
    - Перед единицами измерения
    НЕ трогает титульник.
    """
    import re
    count = 0
    NBSP = '\u00A0'

    def process_text(text):
        if not text:
            return text, False

        new_text = text
        original = text

        # После №, §
        new_text = re.sub(r'([№§])\s+', r'\1' + NBSP, new_text)

        # После сокращений
        new_text = re.sub(r'\b(рис|табл|гл|стр|см|пп?|гг?|др|руб|коп|тыс|млн|млрд)\.\s+',
                         r'\1.' + NBSP, new_text, flags=re.IGNORECASE)

        # Инициалы
        new_text = re.sub(r'([А-ЯA-Z])\.\s+([А-ЯA-Z])\.\s+([А-ЯA-Z])',
                         r'\1.' + NBSP + r'\2.' + NBSP + r'\3', new_text)
        new_text = re.sub(r'([А-ЯA-Z])\.\s+([А-ЯA-Z][а-яa-z])',
                         r'\1.' + NBSP + r'\2', new_text)

        # Единицы измерения
        units = r'(кг|г|мг|т|км|м|см|мм|л|мл|шт|ч|мин|сек|с|руб|коп|%|°C|К|Вт|кВт|В|А|Гц|байт|Кб|Мб|Гб|Тб)'
        new_text = re.sub(r'(\d)\s+' + units + r'\b', r'\1' + NBSP + r'\2', new_text)

        # Годы
        new_text = re.sub(r'(\d{4})\s+(г|гг)\.', r'\1' + NBSP + r'\2.', new_text)

        # Даты
        months = r'(января|февраля|марта|апреля|мая|июня|июля|августа|сентября|октября|ноября|декабря)'
        new_text = re.sub(r'(\d{1,2})\s+' + months, r'\1' + NBSP + r'\2', new_text, flags=re.IGNORECASE)

        return new_text, new_text != original

    for i, para in enumerate(doc.paragraphs):
        if i < title_end_idx:
            continue
        for run in para.runs:
            if run.text:
                new_text, changed = process_text(run.text)
                if changed:
                    run.text = new_text
                    count += 1

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        if run.text:
                            new_text, changed = process_text(run.text)
                            if changed:
                                run.text = new_text
                                count += 1

    logger.info(f"Добавлено неразрывных пробелов в {count} runs")


def _fix_quotes(doc: Document, title_end_idx: int) -> None:
    """
    Заменяет все типы кавычек на русские «ёлочки».
    " " " ' → «»
    НЕ трогает титульник.
    """
    count = 0

    def process_text(text):
        """Заменяет кавычки на ёлочки с учётом контекста"""
        if not text:
            return text, False

        new_text = text
        changed = False

        # Заменяем парные кавычки
        # Открывающие кавычки (после пробела, начала строки, или открывающей скобки)
        import re

        # Паттерн для открывающей кавычки: начало строки, пробел, или ( перед кавычкой
        # Заменяем " " на « если это открывающая позиция

        result = []
        i = 0
        in_quote = False

        while i < len(new_text):
            char = new_text[i]

            if char in '""„"\'':
                # Определяем, открывающая или закрывающая
                prev_char = new_text[i-1] if i > 0 else ' '
                next_char = new_text[i+1] if i < len(new_text) - 1 else ' '

                # Открывающая: после пробела, начала, (, [, или другой открывающей кавычки
                is_opening = prev_char in ' \t\n([«' or i == 0

                if is_opening:
                    result.append('«')
                    in_quote = True
                else:
                    result.append('»')
                    in_quote = False
                changed = True
            elif char == '«' or char == '»':
                result.append(char)
            else:
                result.append(char)

            i += 1

        return ''.join(result), changed

    for i, para in enumerate(doc.paragraphs):
        if i < title_end_idx:
            continue
        for run in para.runs:
            if not run.text:
                continue

            new_text, changed = process_text(run.text)
            if changed:
                run.text = new_text
                count += 1

    # Также в таблицах (они не в титульнике)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        if not run.text:
                            continue

                        new_text, changed = process_text(run.text)
                        if changed:
                            run.text = new_text
                            count += 1

    logger.info(f"Исправлено кавычек в {count} runs")


def _remove_empty_paragraphs(doc: Document, title_end_idx: int) -> None:
    """
    Удаляет лишние пустые параграфы (оставляет максимум 1 пустой подряд).
    Не трогает титульник.
    НЕ удаляет параграфы с рисунками!
    """
    # Собираем индексы параграфов для удаления
    to_remove = []
    consecutive_empty = 0

    for i, para in enumerate(doc.paragraphs):
        if i < title_end_idx:
            continue

        # Проверяем: пустой ли параграф И нет ли в нём рисунка
        is_empty_text = not para.text.strip()
        has_image = _has_image(para)

        if is_empty_text and not has_image:
            consecutive_empty += 1
            if consecutive_empty > 1:
                to_remove.append(i)
        else:
            consecutive_empty = 0

    # Удаляем с конца, чтобы не сбить индексы
    for i in reversed(to_remove):
        para = doc.paragraphs[i]
        p = para._element
        p.getparent().remove(p)

    logger.info(f"Удалено {len(to_remove)} лишних пустых параграфов")


def _fix_colons(doc: Document, title_end_idx: int) -> None:
    """
    Исправляет форматирование двоеточий:
    1. Убирает пробел перед двоеточием: "слово :" -> "слово:"
    2. После двоеточия слово с маленькой буквы: "Тема: Слово" -> "Тема: слово"
    НЕ трогает титульник.
    """
    import re
    count = 0

    def process_text(text):
        new_text = text

        # 1. Убираем пробел перед двоеточием
        new_text = new_text.replace(' :', ':')
        new_text = new_text.replace('\u00A0:', ':')  # неразрывный пробел

        # 2. После двоеточия с пробелом - делаем букву строчной
        def lowercase_after_colon(match):
            colon_space = match.group(1)  # ": "
            letter = match.group(2)  # заглавная буква
            rest = match.group(3)  # остаток слова

            # Не трогаем если это аббревиатура (следующая буква тоже заглавная)
            if rest and len(rest) > 0 and rest[0].isupper():
                return match.group(0)  # оставляем как есть (например API, HTTP)

            return colon_space + letter.lower() + rest

        new_text = re.sub(r'(:\s)([А-ЯA-Z])([а-яa-zА-ЯA-Z]*)', lowercase_after_colon, new_text)

        return new_text

    def process_paragraph_runs(runs):
        """Обрабатывает runs параграфа, учитывая переходы между runs"""
        nonlocal count

        for i, run in enumerate(runs):
            if not run.text:
                continue

            original = run.text
            new_text = process_text(original)

            if new_text != original:
                run.text = new_text
                count += 1

        # Обрабатываем переходы между runs
        # Если run заканчивается на ": " или ":\n", а следующий начинается с заглавной
        for i in range(len(runs) - 1):
            curr_run = runs[i]
            next_run = runs[i + 1]

            if not curr_run.text or not next_run.text:
                continue

            curr_text = curr_run.text
            next_text = next_run.text

            # Проверяем заканчивается ли текущий run на двоеточие с пробелом
            if curr_text.rstrip().endswith(':') or curr_text.endswith(': '):
                # Проверяем начинается ли следующий run с заглавной буквы
                stripped_next = next_text.lstrip()
                if stripped_next and stripped_next[0].isupper():
                    # Проверяем что это не аббревиатура
                    if len(stripped_next) > 1 and not stripped_next[1].isupper():
                        # Делаем первую букву строчной
                        leading_spaces = len(next_text) - len(next_text.lstrip())
                        next_run.text = next_text[:leading_spaces] + stripped_next[0].lower() + stripped_next[1:]
                        count += 1

    for i, para in enumerate(doc.paragraphs):
        if i < title_end_idx:
            continue
        process_paragraph_runs(list(para.runs))

    # Также обрабатываем таблицы (они не в титульнике)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    process_paragraph_runs(list(para.runs))

    logger.info(f"Исправлено двоеточий в {count} runs")


def _fix_dashes(doc: Document, title_end_idx: int) -> None:
    """
    Заменяет длинные тире (—) и дефисы с пробелами на короткие тире (–).
    Обрабатывает случаи когда дефис находится в отдельном run.
    НЕ трогает титульник.
    """
    count = 0

    def process_runs(runs):
        nonlocal count

        for i, run in enumerate(runs):
            if not run.text:
                continue

            original = run.text
            new_text = original

            # Заменяем длинное тире на короткое
            new_text = new_text.replace('—', '–')

            # Заменяем дефис с пробелами на короткое тире
            new_text = new_text.replace(' - ', ' – ')
            new_text = new_text.replace('\u00A0-\u00A0', ' – ')
            new_text = new_text.replace('\u00A0- ', ' – ')
            new_text = new_text.replace(' -\u00A0', ' – ')

            if new_text != original:
                run.text = new_text
                count += 1

        # Обрабатываем случай когда дефис в отдельном run: "слово " + "- " + "слово"
        # или "слово " + "-" + " слово"
        for i in range(len(runs)):
            run = runs[i]
            if not run.text:
                continue

            text = run.text

            # Случай 1: run содержит только "- " или "-"
            if text.strip() == '-':
                # Проверяем что предыдущий run заканчивается на пробел
                # и следующий run начинается с пробела или это "- "
                prev_ok = (i > 0 and runs[i-1].text and runs[i-1].text.endswith(' '))
                next_ok = (i < len(runs)-1 and runs[i+1].text and
                          (runs[i+1].text.startswith(' ') or text == '- '))

                if prev_ok or text == '- ' or text == ' -':
                    run.text = text.replace('-', '–')
                    count += 1

            # Случай 2: run это "- " (дефис с пробелом)
            elif text == '- ':
                run.text = '– '
                count += 1

            # Случай 3: run это " -" (пробел с дефисом)
            elif text == ' -':
                run.text = ' –'
                count += 1

    for i, para in enumerate(doc.paragraphs):
        if i < title_end_idx:
            continue
        process_runs(list(para.runs))

    # Также проверяем таблицы (они не в титульнике)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    process_runs(list(para.runs))

    logger.info(f"Заменено тире в {count} runs")


def _fix_numbering_styles(doc: Document) -> None:
    """
    Исправляет стили нумерации в numbering.xml.
    Устанавливает Times New Roman 14pt для всех номеров списков.
    """
    # Получаем доступ к numbering part
    try:
        numbering_part = doc.part.numbering_part
        if numbering_part is None:
            return

        numbering_elm = numbering_part._element

        # Пространства имён
        nsmap = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}

        # Проходим по всем abstractNum
        for abstractNum in numbering_elm.findall('.//w:abstractNum', nsmap):
            for lvl in abstractNum.findall('.//w:lvl', nsmap):
                # Ищем или создаём rPr
                rPr = lvl.find('.//w:rPr', nsmap)

                if rPr is None:
                    # Создаём rPr
                    rPr = OxmlElement('w:rPr')
                    lvl.append(rPr)

                # Удаляем старые настройки шрифта
                for old_elem in rPr.findall('.//w:rFonts', nsmap):
                    rPr.remove(old_elem)
                for old_elem in rPr.findall('.//w:sz', nsmap):
                    rPr.remove(old_elem)
                for old_elem in rPr.findall('.//w:szCs', nsmap):
                    rPr.remove(old_elem)

                # Добавляем Times New Roman
                rFonts = OxmlElement('w:rFonts')
                rFonts.set(qn('w:ascii'), settings.GOST_FONT_NAME)
                rFonts.set(qn('w:hAnsi'), settings.GOST_FONT_NAME)
                rFonts.set(qn('w:cs'), settings.GOST_FONT_NAME)
                rFonts.set(qn('w:eastAsia'), settings.GOST_FONT_NAME)
                rPr.append(rFonts)

                # Добавляем размер 14pt (28 half-points)
                sz = OxmlElement('w:sz')
                sz.set(qn('w:val'), str(settings.GOST_FONT_SIZE * 2))
                rPr.append(sz)

                szCs = OxmlElement('w:szCs')
                szCs.set(qn('w:val'), str(settings.GOST_FONT_SIZE * 2))
                rPr.append(szCs)

        logger.info("Стили нумерации исправлены")
    except Exception as e:
        logger.warning(f"Не удалось исправить стили нумерации: {e}")


def _fix_list_punctuation(doc: Document, title_end_idx: int) -> None:
    """
    Исправляет знаки препинания в нумерованных списках:
    - Все элементы кроме последнего заканчиваются на ;
    - Последний элемент заканчивается на .

    НЕ трогает библиографический список.
    """
    # Находим начало библиографии, чтобы не трогать её
    bibliography_start_idx = None
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip().upper()
        if 'БИБЛИОГРАФИЧЕСКИЙ' in text or 'СПИСОК ЛИТЕРАТУРЫ' in text:
            bibliography_start_idx = i
            break

    # Собираем все элементы списков с их индексами
    list_items = []

    for i, para in enumerate(doc.paragraphs):
        if i < title_end_idx:
            continue

        # Пропускаем библиографию
        if bibliography_start_idx and i >= bibliography_start_idx:
            continue

        if _is_list_item(para) and para.text.strip():
            list_items.append(i)

    if not list_items:
        return

    # Группируем последовательные элементы в отдельные списки
    lists = []
    current_list = [list_items[0]]

    for i in range(1, len(list_items)):
        # Проверяем, является ли это продолжением списка
        # (индексы идут подряд или с небольшим разрывом на пустые параграфы)
        prev_idx = list_items[i - 1]
        curr_idx = list_items[i]

        # Проверяем, есть ли между ними только пустые параграфы
        is_continuation = True
        for j in range(prev_idx + 1, curr_idx):
            if doc.paragraphs[j].text.strip():
                # Между элементами есть непустой параграф - это новый список
                is_continuation = False
                break

        if is_continuation and (curr_idx - prev_idx) <= 3:
            current_list.append(curr_idx)
        else:
            lists.append(current_list)
            current_list = [curr_idx]

    lists.append(current_list)

    # Обрабатываем каждый список
    for lst in lists:
        if len(lst) < 2:
            # Список из одного элемента - ставим точку
            para = doc.paragraphs[lst[0]]
            _set_paragraph_ending(para, '.')
        else:
            # Все кроме последнего - точка с запятой
            for idx in lst[:-1]:
                para = doc.paragraphs[idx]
                _set_paragraph_ending(para, ';')

            # Последний - точка
            para = doc.paragraphs[lst[-1]]
            _set_paragraph_ending(para, '.')

    logger.info(f"Обработано {len(lists)} списков, {len(list_items)} элементов")


def _set_paragraph_ending(paragraph, ending: str) -> None:
    """
    Устанавливает правильное окончание параграфа.
    Работает и с обычным текстом, и с гиперссылками.
    """
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    text = paragraph.text.rstrip()

    if not text:
        return

    # Если уже заканчивается на нужный знак - не трогаем
    if text.endswith(ending):
        return

    # Проверяем, есть ли гиперссылки в параграфе
    hyperlinks = paragraph._element.findall('.//w:hyperlink',
                                            {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})

    if hyperlinks:
        # Есть гиперссылки - нужно добавить run после последней гиперссылки
        # Создаём новый run элемент
        new_run = OxmlElement('w:r')

        # Добавляем свойства шрифта
        rPr = OxmlElement('w:rPr')
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:ascii'), 'Times New Roman')
        rFonts.set(qn('w:hAnsi'), 'Times New Roman')
        rPr.append(rFonts)

        sz = OxmlElement('w:sz')
        sz.set(qn('w:val'), '28')  # 14pt = 28 half-points
        rPr.append(sz)

        new_run.append(rPr)

        # Добавляем текст
        t = OxmlElement('w:t')
        t.text = ending
        new_run.append(t)

        # Добавляем run в конец параграфа (после всех элементов)
        paragraph._element.append(new_run)
        return

    # Нет гиперссылок - работаем с runs
    if paragraph.runs:
        last_run = None
        for run in reversed(paragraph.runs):
            if run.text and run.text.strip():
                last_run = run
                break

        if last_run:
            run_text = last_run.text.rstrip()
            if run_text and run_text[-1] in '.;,:!':
                last_run.text = run_text[:-1] + ending
            else:
                last_run.text = run_text + ending
            return

    # Fallback - добавляем новый run
    new_run = paragraph.add_run(ending)
    new_run.font.name = 'Times New Roman'
    new_run.font.size = Pt(14)


def _add_page_break_before(paragraph) -> None:
    """Добавляет разрыв страницы перед параграфом через pageBreakBefore"""
    p = paragraph._element
    pPr = p.get_or_add_pPr()

    # Удаляем существующий если есть
    existing = pPr.find(qn('w:pageBreakBefore'))
    if existing is not None:
        pPr.remove(existing)

    # Добавляем новый
    page_break = OxmlElement('w:pageBreakBefore')
    page_break.set(qn('w:val'), 'true')
    pPr.insert(0, page_break)


def _remove_page_break_before(paragraph) -> None:
    """Удаляет разрыв страницы перед параграфом"""
    pPr = paragraph._element.find(qn('w:pPr'))
    if pPr is not None:
        pb = pPr.find(qn('w:pageBreakBefore'))
        if pb is not None:
            pPr.remove(pb)


def format_document(input_file: str, output_file: str = None) -> bool:
    """
    Основная функция для форматирования документа по ГОСТ
    """
    if not output_file:
        base, ext = os.path.splitext(input_file)
        output_file = f"{base}_formatted{ext}"

    if not os.path.exists(input_file):
        logger.error(f"Файл '{input_file}' не найден!")
        return False

    try:
        logger.info("Загрузка документа...")
        doc = Document(input_file)
        logger.info("Документ успешно загружен")

        # Определяем конец титульного листа
        title_end_idx = _find_title_page_end(doc)
        logger.info(f"Титульный лист заканчивается на параграфе {title_end_idx}")

        # НЕ меняем поля страницы - это ломает титульник
        # (документ имеет одну секцию, поля применяются ко всему)
        # Если нужно изменить поля - пользователь делает это вручную
        logger.info("Поля страницы не изменяются (сохраняем оригинальные)")

        # НЕ меняем нумерацию страниц - это может затрагивать титульник
        # Нумерация должна быть настроена в исходном документе
        logger.info("Нумерация страниц не изменяется (сохраняем оригинальную)")

        # НЕ меняем стили нумерации списков - это глобальные стили,
        # они могут затронуть титульник
        logger.info("Стили нумерации списков не изменяются (сохраняем оригинальные)")

        # Обработка подписей к рисункам
        logger.info("Обработка подписей к рисункам...")
        _process_figures(doc, title_end_idx)

        # Удаление пустых параграфов перед главными заголовками
        logger.info("Обработка разрывов страниц...")
        _process_page_breaks(doc, title_end_idx)

        # Форматирование документа (кроме титульника)
        logger.info("Форматирование документа...")
        _format_document_content(doc, title_end_idx)

        # Исправление знаков препинания в списках
        logger.info("Исправление знаков препинания в списках...")
        _fix_list_punctuation(doc, title_end_idx)

        # Замена длинных тире на средние
        logger.info("Замена длинных тире на средние...")
        _fix_dashes(doc, title_end_idx)

        # Исправление двоеточий
        logger.info("Исправление двоеточий...")
        _fix_colons(doc, title_end_idx)

        # Убираем множественные пробелы
        logger.info("Убираем множественные пробелы...")
        _fix_multiple_spaces(doc, title_end_idx)

        # Унификация кавычек
        logger.info("Унификация кавычек...")
        _fix_quotes(doc, title_end_idx)

        # Неразрывные пробелы
        logger.info("Добавление неразрывных пробелов...")
        _fix_non_breaking_spaces(doc, title_end_idx)

        # Форматирование таблиц
        logger.info("Форматирование таблиц...")
        _format_tables(doc)

        # Подписи таблиц
        logger.info("Проверка подписей таблиц...")
        _add_table_captions(doc)

        # Замена сокращений
        logger.info("Замена сокращений...")
        _fix_abbreviations(doc, title_end_idx)

        # Форматирование библиографии
        logger.info("Форматирование библиографии...")
        _format_bibliography(doc)

        # Удаление лишних пустых строк
        logger.info("Удаление лишних пустых строк...")
        _remove_empty_paragraphs(doc, title_end_idx)

        # Сохранение результата
        logger.info("Сохранение документа...")
        doc.save(output_file)
        logger.info(f"Документ успешно сохранен: {output_file}")

        return True

    except Exception as e:
        logger.error(f"Ошибка при обработке документа: {str(e)}")
        logger.debug(f"Детали ошибки: {traceback.format_exc()}")
        return False


def _process_figures(doc: Document, title_end_idx: int) -> None:
    """
    Обрабатывает рисунки: добавляет недостающие подписи и исправляет нумерацию.
    """
    figure_number = 0
    i = title_end_idx

    while i < len(doc.paragraphs):
        para = doc.paragraphs[i]

        # Проверяем, есть ли изображение в параграфе
        if _has_image(para):
            figure_number += 1

            # Проверяем следующий параграф на наличие подписи
            next_para = doc.paragraphs[i + 1] if i + 1 < len(doc.paragraphs) else None

            if next_para:
                next_text = next_para.text.strip()

                if _is_figure_caption(next_text):
                    # Подпись есть — проверяем/исправляем номер
                    _fix_figure_caption_number(next_para, figure_number)
                else:
                    # Подписи нет — добавляем
                    logger.info(f"Добавляем подпись для рисунка {figure_number}")
                    _insert_figure_caption(doc, i + 1, figure_number)

        i += 1


def _fix_figure_caption_number(paragraph, correct_number: int) -> None:
    """
    Исправляет номер в подписи рисунка.
    'Рисунок  - описание' -> 'Рисунок 1 – описание'
    'Рисунок 5 - описание' -> 'Рисунок 1 – описание'
    """
    text = paragraph.text.strip()

    # Паттерн: Рисунок [номер] [-–—] описание
    match = re.match(r'^(Рисунок)\s*(\d*)\s*[-–—]\s*(.*)$', text, re.IGNORECASE)

    if match:
        prefix = match.group(1)  # "Рисунок"
        description = match.group(3)  # описание после тире

        # Используем короткое тире –
        new_text = f"{prefix} {correct_number} – {description}"

        # Очищаем параграф и записываем новый текст
        paragraph.clear()
        run = paragraph.add_run(new_text)
        run.font.name = settings.GOST_FONT_NAME
        run.font.size = Pt(settings.GOST_FONT_SIZE)
        run._element.rPr.rFonts.set(qn('w:eastAsia'), settings.GOST_FONT_NAME)


def _insert_figure_caption(doc: Document, insert_index: int, figure_number: int) -> None:
    """
    Вставляет новую подпись к рисунку.
    """
    # Получаем параграф, ПЕРЕД которым нужно вставить
    if insert_index < len(doc.paragraphs):
        target_para = doc.paragraphs[insert_index]
        # Вставляем новый параграф перед target (используем короткое тире –)
        new_para = target_para.insert_paragraph_before(f"Рисунок {figure_number} – ")
    else:
        # Добавляем в конец
        new_para = doc.add_paragraph(f"Рисунок {figure_number} – ")

    # Форматируем
    new_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    new_para.paragraph_format.first_line_indent = Cm(0)
    new_para.paragraph_format.space_before = Pt(0)
    new_para.paragraph_format.space_after = Pt(0)
    new_para.paragraph_format.line_spacing = settings.GOST_LINE_SPACING

    for run in new_para.runs:
        run.font.name = settings.GOST_FONT_NAME
        run.font.size = Pt(settings.GOST_FONT_SIZE)


def _process_page_breaks(doc: Document, title_end_idx: int) -> None:
    """
    Обрабатывает разрывы страниц:
    - Удаляет пустые параграфы перед главными заголовками
    - Добавляет корректные разрывы страниц
    """
    # Собираем индексы главных заголовков
    main_heading_indices = []

    for i, para in enumerate(doc.paragraphs):
        if i < title_end_idx:
            continue
        text = para.text.strip()
        if text and _is_main_heading(text):
            main_heading_indices.append(i)

    # Для каждого главного заголовка удаляем пустые параграфы перед ним
    # и добавляем pageBreakBefore
    # ВАЖНО: идём с конца, чтобы не сбить индексы

    paragraphs_to_remove = []

    for heading_idx in main_heading_indices:
        # Ищем пустые параграфы перед заголовком
        j = heading_idx - 1
        while j >= title_end_idx:
            prev_para = doc.paragraphs[j]
            if not prev_para.text.strip():
                paragraphs_to_remove.append(j)
                j -= 1
            else:
                break

    # Удаляем пустые параграфы (с конца, чтобы не сбить индексы)
    for idx in sorted(set(paragraphs_to_remove), reverse=True):
        para = doc.paragraphs[idx]
        p = para._element
        p.getparent().remove(p)

    logger.info(f"Удалено {len(paragraphs_to_remove)} пустых параграфов перед заголовками")

    # Теперь добавляем разрывы страниц к главным заголовкам
    # (индексы могли измениться, ищем заново)
    for i, para in enumerate(doc.paragraphs):
        if i < title_end_idx:
            continue
        text = para.text.strip()
        if text and _is_main_heading(text):
            _add_page_break_before(para)


def _format_document_content(doc: Document, title_end_idx: int) -> None:
    """
    Форматирует содержимое документа, пропуская титульный лист.
    """
    for i, paragraph in enumerate(doc.paragraphs):
        if i < title_end_idx:
            continue

        text = paragraph.text.strip()
        if not text:
            continue

        if _is_main_heading(text):
            _format_main_heading(paragraph)
        elif _is_subheading(text):
            _format_subheading(paragraph)
        elif _is_figure_caption(text) or _is_table_caption(text):
            _format_caption(paragraph)
        elif _is_list_item(paragraph):
            _format_list_item(paragraph)
        else:
            _format_regular_paragraph(paragraph)


def _format_main_heading(paragraph) -> None:
    """
    Форматирует главный заголовок: 16pt BOLD, CENTER, без отступа
    """
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    paragraph.paragraph_format.first_line_indent = Cm(0)
    paragraph.paragraph_format.left_indent = Cm(0)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = settings.GOST_LINE_SPACING

    for run in paragraph.runs:
        run.font.name = settings.GOST_FONT_NAME
        run.font.size = Pt(16)
        run.font.bold = True
        run._element.rPr.rFonts.set(qn('w:eastAsia'), settings.GOST_FONT_NAME)


def _format_subheading(paragraph) -> None:
    """
    Форматирует подзаголовок: 14pt BOLD, CENTER, без отступа
    """
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    paragraph.paragraph_format.first_line_indent = Cm(0)
    paragraph.paragraph_format.left_indent = Cm(0)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = settings.GOST_LINE_SPACING

    for run in paragraph.runs:
        run.font.name = settings.GOST_FONT_NAME
        run.font.size = Pt(settings.GOST_FONT_SIZE)
        run.font.bold = True
        run._element.rPr.rFonts.set(qn('w:eastAsia'), settings.GOST_FONT_NAME)


def _format_caption(paragraph) -> None:
    """
    Форматирует подпись к рисунку/таблице: 14pt, CENTER, без отступа
    """
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    paragraph.paragraph_format.first_line_indent = Cm(0)
    paragraph.paragraph_format.left_indent = Cm(0)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = settings.GOST_LINE_SPACING

    for run in paragraph.runs:
        run.font.name = settings.GOST_FONT_NAME
        run.font.size = Pt(settings.GOST_FONT_SIZE)
        run.font.bold = False
        run._element.rPr.rFonts.set(qn('w:eastAsia'), settings.GOST_FONT_NAME)


def _format_list_item(paragraph) -> None:
    """
    Форматирует элемент списка:
    - 14pt Times New Roman
    - По ширине (JUSTIFY)
    - Первая строка (с номером) с отступом 1.25 см
    - Последующие строки от начала (0 см)
    """
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Отступы:
    # - first_line_indent = 1.25 см (красная строка с номером)
    # - left_indent = 0 (последующие строки от начала)
    paragraph.paragraph_format.left_indent = Cm(0)
    paragraph.paragraph_format.first_line_indent = Cm(settings.GOST_INDENT_CM)  # 1.25 см

    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = settings.GOST_LINE_SPACING

    # Форматируем шрифт
    for run in paragraph.runs:
        run.font.name = settings.GOST_FONT_NAME
        run.font.size = Pt(settings.GOST_FONT_SIZE)
        run._element.rPr.rFonts.set(qn('w:eastAsia'), settings.GOST_FONT_NAME)


def _format_regular_paragraph(paragraph) -> None:
    """
    Форматирует обычный абзац: 14pt, JUSTIFY, красная строка 1.25 см
    """
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    paragraph.paragraph_format.first_line_indent = Cm(settings.GOST_INDENT_CM)
    paragraph.paragraph_format.left_indent = Cm(0)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = settings.GOST_LINE_SPACING

    for run in paragraph.runs:
        run.font.name = settings.GOST_FONT_NAME
        run.font.size = Pt(settings.GOST_FONT_SIZE)
        run._element.rPr.rFonts.set(qn('w:eastAsia'), settings.GOST_FONT_NAME)


def _add_page_numbers(doc: Document) -> None:
    """
    Добавляет нумерацию страниц.
    Титульник = страница 1, но номер не отображается.
    """
    for section in doc.sections:
        section.different_first_page_header_footer = True

        footer = section.footer
        footer.is_linked_to_previous = False

        for para in footer.paragraphs:
            para.clear()

        paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        run = paragraph.add_run()
        fld_char_begin = OxmlElement('w:fldChar')
        fld_char_begin.set(qn('w:fldCharType'), 'begin')

        instr_text = OxmlElement('w:instrText')
        instr_text.set(qn('xml:space'), 'preserve')
        instr_text.text = 'PAGE'

        fld_char_end = OxmlElement('w:fldChar')
        fld_char_end.set(qn('w:fldCharType'), 'end')

        run._r.append(fld_char_begin)
        run._r.append(instr_text)
        run._r.append(fld_char_end)

        run.font.name = settings.GOST_FONT_NAME
        run.font.size = Pt(settings.GOST_FONT_SIZE)